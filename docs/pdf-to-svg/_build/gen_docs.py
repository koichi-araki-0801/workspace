# -*- coding: utf-8 -*-
"""PdfToSvg の 2 文書（エンジニア向け設計書 / 事務担当者向け操作手順書）を Word(.docx) で生成する。

- 出力先: docs/pdf-to-svg/PdfToSvg_設計書.docx, PdfToSvg_操作手順書.docx
- 画面キャプチャ: docs/pdf-to-svg/images/*.png（capture_screens.py が事前生成）
- 依存: python-docx（成果物生成のための一時依存。ツール本体には追加しない）
"""
from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# リポジトリルートからの相対解決 (このファイルは <repo>/docs/pdf-to-svg/_build/ にある)。
DOCS = pathlib.Path(__file__).resolve().parents[1]
IMG = DOCS / "images"

JP = "Meiryo UI"          # 本文用フォント（CLAUDE.md「Word ドキュメント」規約: 本文は Meiryo UI 統一）
MONO = "ＭＳ ゴシック"      # 等幅（ASCII図 + 日本語）。Meiryo UI は非等幅のため図のけた揃え用に等幅を維持
ACCENT = RGBColor(0x1F, 0x5C, 0x99)
INK = RGBColor(0x20, 0x24, 0x2C)
MUTED = RGBColor(0x60, 0x68, 0x74)


# ---------------------------------------------------------------- helpers
def _set_run_font(run, name=JP, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def new_doc() -> Document:
    doc = Document()
    # 既定フォントを日本語対応へ
    normal = doc.styles["Normal"]
    normal.font.name = JP
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), JP)
    # 余白
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
    return doc


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size=22, bold=True, color=ACCENT)
    if subtitle:
        ps = doc.add_paragraph()
        rs = ps.add_run(subtitle)
        _set_run_font(rs, size=11, color=MUTED)


def h(doc, text, level=1):
    sizes = {1: 15, 2: 12.5, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    _set_run_font(r, size=sizes.get(level, 11), bold=True,
                  color=ACCENT if level == 1 else INK)
    return p


def para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for seg, bold in _segments(text):
        r = p.add_run(seg)
        _set_run_font(r, size=size, bold=bold)
    return p


def _segments(text):
    """**強調** を太字に分解する簡易パーサ。"""
    out = []
    for i, part in enumerate(text.split("**")):
        if part:
            out.append((part, i % 2 == 1))
    return out or [("", False)]


def bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    for seg, bold in _segments(text):
        r = p.add_run(seg)
        _set_run_font(r, size=size, bold=bold)
    return p


def numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for seg, bold in _segments(text):
        r = p.add_run(seg)
        _set_run_font(r, size=size, bold=bold)
    return p


def code(doc, text, size=8.8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    # 薄い網掛け
    shd = p._p.get_or_add_pPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F2F4F7"})
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run(text)
    _set_run_font(r, name=MONO, size=size)


def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        _set_run_font(run, size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_run_font(run, size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def image(doc, name, width_cm=16.0, caption=None):
    path = IMG / name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        _set_run_font(r, size=9, italic=True, color=MUTED)
        c.paragraph_format.space_after = Pt(8)


def callout(doc, text, fill="FFF6E5"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
    pPr.append(shd)
    for seg, bold in _segments(text):
        r = p.add_run(seg)
        _set_run_font(r, size=10, bold=bold)


def page_break(doc):
    doc.add_page_break()


# ================================================================ 設計書
def build_design():
    doc = new_doc()
    title(doc, "PdfToSvg 詳細設計書",
          "PDF → 編集可能 SVG 変換ツール ／ エンジニア向け（保守・改修用）")
    para(doc, "版: 1.0　　対象リビジョン: 最新版　　"
              "本書は実装の構造・依存・ライフサイクルを記述し、引き継ぎと改修を可能にすることを目的とする。",
         size=9.5)

    h(doc, "1. 概要・目的")
    para(doc, "PdfToSvg は、PDF を**編集可能な SVG** へ変換する Windows デスクトップツールである。"
              "非エンジニア（事務担当者）が単独で次の作業を行えることを狙いとする。")
    bullet(doc, "**表ヘッダ等の用語統一・言い換え**: JSON 永続辞書による自動置換（全角/半角差を吸収）。")
    bullet(doc, "**不要領域のトリミング**: 要素単位の削除・矩形クロップ・枠線追加。")
    para(doc, "出力 SVG はテキストを**本物の `<text>` 要素**として保持するため、後から検索・再編集できる。"
              "使用したグリフのみを WOFF2 サブセットで埋め込むため出力は軽量。")
    para(doc, "**配布形態**: PyInstaller の onedir 形式によるポータブル exe（インストーラ不要・"
              "レジストリ非使用）。フォルダごとコピーで別端末へ展開可能。")

    h(doc, "2. 動作環境")
    table(doc, ["項目", "内容"], [
        ["OS", "Windows 10 / 11"],
        ["UI 表示", "システム標準の Microsoft Edge（Chromium）を「アプリモード」で使用"],
        ["ランタイム同梱", "Python 3.10+ ランタイム（exe に同梱）。Qt / PySide6 は非同梱"],
        ["ネットワーク", "127.0.0.1 のローカル通信のみ。外部通信なし"],
        ["権限", "管理者権限不要。ユーザー権限で動作"],
    ], widths=[4.0, 12.0])
    para(doc, "Edge をアプリモードで開く方式のため QtWebEngine ランタイムを同梱せずに済み、"
              "配布物が 150–200MB から ~50MB へ縮小されている。")

    h(doc, "3. アーキテクチャ")
    para(doc, "設計原則は**「モデルが真実（model is truth）」**。編集は Web UI から `/rpc` 経由で"
              "コマンドを Undo スタックへ push してモデルを更新し、SVG はモデルから直接書き出す"
              "（決定的・GUI 非依存でユニットテスト可能）。")
    code(doc,
         "PDF\n"
         " └→ engine (PyMuPDF抽出 + vector/scan判定)\n"
         "      └→ model (Document / Page / Element)\n"
         "           └→ dictionary (NFKC正規化 + ヘッダ検出 + 自動適用)\n"
         "                └→ export (model→SVG 決定的シリアライズ) ……出力 & 画面プレビュー\n"
         "\n"
         "          ↑↓  web.server (127.0.0.1 HTTP  /rpc・/upload)\n"
         "          ←→  resources/web (Edge アプリ窓の UI)")
    para(doc, "代表的な操作の流れ（用語置換）:")
    numbered(doc, "UI が `/upload` へ PDF バイト列を送信 → `loader` が `Document` を構築し、辞書を自動適用。")
    numbered(doc, "UI が `pageSvg` RPC を呼び、`page_to_svg(annotate=True)` がモデルから SVG を生成して表示。")
    numbered(doc, "削除/クロップ等の編集は対応コマンドを Undo スタックへ push してモデルを更新。")
    numbered(doc, "`undo`/`redo` RPC でコマンドの undo()/redo() を実行し、SVG を再生成。"
                  "辞書適用はマクロ化され「1 操作 = 1 Undo ステップ」。")

    h(doc, "4. モジュール構成")
    para(doc, "ディレクトリと主な責務:")
    table(doc, ["パス", "責務"], [
        ["src/app.py", "エントリポイント。Edge 探索・アプリ窓起動・ライフサイクル管理"],
        ["src/config.py", "パス解決（frozen exe / ソース実行の双方）。data/・fonts/・resources/ の基準"],
        ["src/engine/pdf_engine.py", "PyMuPDF で PDF を抽出し Document モデルへ"],
        ["src/engine/classify.py", "ベクター / スキャン（画像主体）ページの判定"],
        ["src/model/document.py, elements.py", "Document / Page / 各 Element（Text/Line/Rect/Path/Image）"],
        ["src/model/fonts.py", "PDF フォント名 → Windows 標準 / 同梱フォントへのマッピング、ウェイト保持"],
        ["src/export/svg_exporter.py", "モデル → SVG 文字列の決定的シリアライズ。crop_rect を viewBox に反映"],
        ["src/export/font_embed.py", "fontTools によるサブセット化 + base64 WOFF2 @font-face 埋め込み"],
        ["src/dictionary/store.py", "辞書の JSON 永続化（add/upsert/update/delete/all/lookup/import/export）"],
        ["src/dictionary/apply.py", "ヘッダ検出・置換適用・折返しヘッダの畳み込み"],
        ["src/dictionary/normalize.py", "NFKC 正規化（全角/半角差の吸収）"],
        ["src/web/server.py", "ローカル HTTP サーバ（静的配信 + /rpc + /upload + /quit + /ping）"],
        ["src/web/rpc_methods.py", "RPC 実体。WebSession が複数 PDF・辞書・Undo を保持"],
        ["src/web/commands.py, undo_stack.py", "コマンドパターン（Delete/Crop/ReplaceText/AddElement）と純 Python Undo"],
        ["src/web/loader.py", "PDF バイト列 → Document（一時ファイル経由）"],
        ["resources/web/", "Edge アプリ窓の UI（index.html / app.js / rpc.js / styles.css）"],
        ["fonts/", "同梱フォント（BIZ UDPゴシック / Noto Serif JP, SIL OFL）"],
        ["packaging/pdftosvg.spec", "PyInstaller 設定（Qt 除外・fonts/resources 同梱）"],
        ["tests/", "pytest（engine/model/web、決定的 SVG の byte 検証）"],
    ], widths=[6.2, 9.8])

    h(doc, "4.1 HTTP エンドポイント", level=2)
    table(doc, ["メソッド/パス", "役割"], [
        ["GET /  (静的)", "resources/web 配下の配信（拡張子ホワイトリスト + パストラバーサル防止）"],
        ["POST /rpc", "{method, args} を dispatch し {ok, data} を返す JSON-RPC"],
        ["POST /upload?name=", "PDF バイト列を受領し Document 化、辞書を自動適用"],
        ["POST /quit", "ウィンドウ閉鎖ビーコン（quit_event を立てる）"],
        ["POST /ping", "生存ハートビート（idle watchdog の last_seen 更新）"],
    ], widths=[4.6, 11.4])
    para(doc, "ThreadingHTTPServer のため、セッション状態の変更は `server.lock` で直列化する。", size=9.5)

    h(doc, "4.2 主な RPC メソッド", level=2)
    table(doc, ["メソッド", "内容"], [
        ["state", "ファイル/ページ一覧・変更フラグのスナップショット"],
        ["pageSvg", "指定ページの SVG（注釈付き）"],
        ["planPage / removedList", "置換予定一覧 / 削除済み要素一覧"],
        ["applyDelete / deleteRegion / addBorder", "削除・クロップ・枠線（Undo へ push）"],
        ["dictList / dictAdd / reapplyDict", "辞書の参照・追加・全ファイル再適用"],
        ["dictJson / dictImportJson", "辞書の JSON 書き出し / 読み込み"],
        ["setOnlyHeaders", "「ヘッダのみに適用」フラグの切替"],
        ["undo / redo", "Undo / Redo"],
        ["exportSvg", "指定ページの SVG を生成（書き出し）"],
        ["removeFile", "選択ファイルを一覧から除去"],
    ], widths=[6.0, 10.0])

    h(doc, "5. 技術スタック・依存")
    table(doc, ["コンポーネント", "バージョン", "ライセンス", "役割"], [
        ["PyMuPDF (fitz)", "≥1.24", "AGPL", "PDF 抽出・vector/raster 判定"],
        ["Pillow", "≥10.0", "HPND", "画像処理"],
        ["fontTools", "≥4.50", "MIT", "フォントサブセット（TTF→WOFF2）"],
        ["brotli", "≥1.1", "Apache 2.0", "WOFF2 圧縮"],
        ["setuptools", "≥68", "MIT", "パッケージング"],
        ["PyInstaller (dev)", "≥6.20", "GPL+例外", "exe ビルド"],
        ["pytest (dev)", "≥8.0", "MIT", "テスト"],
    ], widths=[3.6, 2.0, 2.6, 7.8])
    para(doc, "フロントエンドはフレームワーク無しの素の JavaScript（最小フットプリント）。"
              "任意で Vitest による最小テストを持つ。")

    h(doc, "6. データ・設定")
    para(doc, "辞書・設定は exe と同じ場所の `data/` に置くポータブル方式（C ドライブや %APPDATA% に非依存）。")
    table(doc, ["ファイル", "役割"], [
        ["data/dictionary.json", "用語マッピング（人手編集・共有可・NFKC 正規化）"],
        ["data/settings.json", "UI 状態・設定（将来拡張）"],
        ["data/startup.log", "起動診断ログ（exe は console=False のため stderr の代替）"],
    ], widths=[4.6, 11.4])
    para(doc, "辞書フォーマット:")
    code(doc,
         '[\n'
         '  {"source": "元の語", "target": "置換後の語", "enabled": true},\n'
         '  {"source": "Old Header", "target": "New Header", "enabled": false}\n'
         ']')
    para(doc, "パス解決（`config.py`）: frozen 時は exe のあるフォルダ、ソース実行時はリポジトリ直下を基準に "
              "data/・fonts/・resources/ を解決する。これにより配布フォルダを任意の場所へコピーできる。")

    h(doc, "7. フォント")
    para(doc, "同梱フォントは BIZ UDPゴシック（ゴシック代替）と Noto Serif JP（明朝代替）。"
              "いずれも SIL OFL 1.1 で再配布可。元 PDF のウェイト（Light/Regular/Medium/Bold）を保持するため、"
              "フォント名から CSS ウェイトを抽出し `<text>` に数値 `font-weight` を出力する。"
              "出力時は使用グリフのみをサブセット化し base64 WOFF2 として @font-face に埋め込む。")

    h(doc, "8. ビルド・配布")
    numbered(doc, "`build.bat` が隔離 venv（`.venv-build`）を作成し必要依存のみインストール。")
    numbered(doc, "`pyinstaller packaging/pdftosvg.spec --clean --noconfirm` で onedir exe を生成。")
    numbered(doc, "成果物 `dist/PdfToSvg/`（exe + _internal + resources + fonts）をフォルダごと配布。")
    para(doc, "spec は PySide6 / Qt DLL を除外し、fonts/ と resources/ を datas として同梱する。"
              "data/ は初回起動時に exe 隣へ生成される。")

    h(doc, "9. 起動ライフサイクル")
    numbered(doc, "`_find_edge()` がレジストリ（App Paths）→既知パスの順に msedge.exe を探索。")
    numbered(doc, "専用 user-data-dir（temp）を作り `--app=URL` で独立プロセスとして起動。")
    numbered(doc, "ローカル HTTP サーバを daemon スレッドで起動（ポートは 0 指定で OS 自動割当）。")
    numbered(doc, "終了契機は窓を閉じた時の `/quit` ビーコン、または `/ping` 途絶を見張る idle watchdog（60 秒）。")
    numbered(doc, "終了時にサーバを畳み、Edge を terminate、temp プロファイルを削除。")
    callout(doc, "**留意**: サーバ寿命は起動した msedge.exe プロセスに縛らない。Edge はコールド時に"
                 "自身を別プロセスへ再起動するため、起動プロセスの早期終了でサーバを畳むと初回が空白になる。"
                 "終了は /quit ＋ watchdog に一本化している。", fill="EAF2FB")

    h(doc, "10. テスト")
    para(doc, "`tests/` に pytest スイート。SVG 出力は決定的（同一 Document → 同一 SVG）であることを前提に、"
              "engine / model / dictionary / export / web を網羅。フロントは任意で Vitest。"
              "実行は `python -m pytest`。")

    h(doc, "11. ライセンス上の注意")
    callout(doc, "本ツールは **PyMuPDF (AGPL)** に依存するため、**社内 / 身内限定の配布**を前提とする。"
                 "外部・商用配布が必要になった場合は、抽出部 `engine/pdf_engine.py` を BSD ライセンスの "
                 "**pypdfium2** 等へ差し替えるか、Artifex の商用ライセンスを取得すること。"
                 "UI は OS 標準 Edge のアプリモードのため Qt 依存は無い。", fill="FDECEC")

    out = DOCS / "PdfToSvg_設計書.docx"
    doc.save(str(out))
    print("saved", out.name)


# ============================================================== 操作手順書
def build_manual():
    doc = new_doc()
    title(doc, "PdfToSvg 操作手順書",
          "PDF を SVG に変換するツール ／ ご担当者さま向け")
    para(doc, "このツールは、PDF をきれいな画像データ（SVG）に変換するためのものです。"
              "むずかしい知識は不要です。画面の **1 → 2 → 3 → 4** の順に進めるだけで完成します。",
         size=10.5)

    h(doc, "1. このツールでできること")
    bullet(doc, "PDF を、文字を残したまま編集できる画像（SVG）に**変換**します。")
    bullet(doc, "表の見出しなどの**用語を、辞書を使って自動でそろえます**（例: 「Header A」→「見出し A」）。")
    bullet(doc, "**いらない部分を消したり、残す範囲を指定したり、枠線を足したり**できます。")
    para(doc, "完成した SVG は、あとから文字を検索・修正できます。ファイルは軽く、印刷や資料作成にも使えます。")

    h(doc, "2. 起動のしかた")
    numbered(doc, "配布されたフォルダの中の **「run.bat」**（または **「PdfToSvg.exe」**）を**ダブルクリック**します。")
    numbered(doc, "少し待つと、**「PdfToSvg」というアプリの画面**が開きます（ふつうの黒いウィンドウは閉じないでください）。")
    callout(doc, "画面が開かないときは、しばらく待ってからもう一度ダブルクリックしてください。"
                 "それでも開かないときは、情報システム担当へご連絡ください。", fill="EAF2FB")

    page_break(doc)
    h(doc, "3. ステップ 1：PDF を選ぶ")
    para(doc, "最初の画面で、変換したい PDF を選びます。複数の PDF をまとめて選ぶこともできます。")
    image(doc, "step1_select.png", caption="ステップ 1：PDF を選ぶ画面")
    numbered(doc, "青い枠の中をクリックするか、**「ファイルを選ぶ…」**ボタンを押します。")
    numbered(doc, "変換したい PDF を選びます（複数選択も可能）。PDF を枠の上に**ドラッグ&ドロップ**しても入れられます。")
    numbered(doc, "選んだ PDF が下の一覧に表示されます。まちがえたら右の **×** で取り消せます。")
    numbered(doc, "右下の **「次へ」** を押します。")
    callout(doc, "**「次へ」が押せないとき**は、まだ PDF が 1 つも選ばれていません。先に PDF を選んでください。")

    page_break(doc)
    h(doc, "4. ステップ 2：用語を置換する")
    para(doc, "辞書に登録した用語を、自動でそろえます。置き換わった箇所をこの画面で確認します。")
    image(doc, "step2_replace.png", caption="ステップ 2：置換された箇所の確認（右側に置換の一覧）")
    bullet(doc, "中央に変換後のページが表示されます。例では「Header A」が「**見出し A**」に置き換わっています。")
    bullet(doc, "右側の **「確認」** タブに、置き換わった内容（**Header A → 見出し A**）が一覧で出ます。")
    bullet(doc, "一覧の行をクリックすると、ページ上の該当箇所が**青く光って**わかります。")
    bullet(doc, "内容を確認したら **「確認しました」**、直す必要がなければ **「スキップ」** を押します。")

    h(doc, "4.1 用語を辞書に追加する", level=2)
    para(doc, "右側の **「辞書」** タブで、新しい用語を登録できます。登録するとすべてのページに自動で反映されます。")
    image(doc, "step2b_dict.png", caption="ステップ 2：辞書タブ（用語の追加と登録一覧）")
    numbered(doc, "**「辞書」** タブを開きます。")
    numbered(doc, "**「元の語」** に置き換えたい言葉、**「置換後」** に新しい言葉を入力します。")
    numbered(doc, "**「＋ 辞書に追加」** を押すと一覧に追加され、ページに反映されます。")
    numbered(doc, "用語を増やしたあとは **「★ 全ファイルに再適用」** を押すと、選んだすべての PDF に反映されます。")
    callout(doc, "**「ヘッダのみに適用」**にチェックがあると、表の見出し部分だけを置き換えます。"
                 "本文も置き換えたいときはチェックを外してください。", fill="EAF2FB")

    h(doc, "4.2 確認していないページが残っているとき", level=2)
    para(doc, "未確認のページがあると、「次へ」を押したときに画面の下に黄色いバーが出ます。"
              "**「最初の未確認へ」** で確認に戻るか、すべて確認不要なら "
              "**「未確認をすべてスキップして進む」** を押します。")

    page_break(doc)
    h(doc, "5. ステップ 3：削除・枠線の編集")
    para(doc, "いらない部分を消したり、残す範囲を決めたり、枠線を足したりできます。")
    image(doc, "step3_edit.png", caption="ステップ 3：削除・クロップ・枠線の編集画面")
    bullet(doc, "**🎯 選択**：消したい部分をクリックして選び、**「🗑 削除」** を押すと消えます。")
    bullet(doc, "**✂ クロップ**：マウスで四角く囲むと、その**内側だけを残します**（外側は書き出し時に消えます）。")
    bullet(doc, "**⬚ 枠線**：色と太さを選んで四角く囲むと、**枠線を追加**できます。")
    bullet(doc, "まちがえたら右側の **「戻す」**、または **Ctrl+Z** でやり直せます。")
    bullet(doc, "確認できたら **「確認しました」**、編集不要なら **「スキップ」** を押し、**「書き出しへ」** に進みます。")

    page_break(doc)
    h(doc, "6. ステップ 4：SVG に書き出す")
    para(doc, "最後に、設定を確認して SVG ファイルとして保存します。")
    image(doc, "step4_export.png", caption="ステップ 4：書き出し画面")
    numbered(doc, "**「書き出す範囲」** を選びます。ふつうは **「全ページ」** のままで大丈夫です。")
    bullet(doc, "「表示中のページのみ」：いま開いている 1 ページだけ。")
    bullet(doc, "「スキップを除く」：スキップしたページを除いて書き出します。")
    bullet(doc, "「ページを指定」：「1-5, 8」のように番号で指定できます。")
    numbered(doc, "右下の緑色の **「SVG に書き出す」** ボタンを押します。")
    numbered(doc, "**保存先のフォルダ**を選ぶ画面が出るので、保存したい場所を選んで決定します。")
    numbered(doc, "完了すると「○個の SVG を書き出しました。」と表示されます。")
    callout(doc, "**ファイル名**は「元のファイル名_p1.svg, _p2.svg …」のように、ページごとに自動で付きます。")

    page_break(doc)
    h(doc, "7. 困ったとき（よくある質問）")
    table(doc, ["こまりごと", "対処"], [
        ["「次へ」が押せない", "ステップ 1 で PDF がまだ選ばれていません。先に PDF を選んでください。"],
        ["黄色い警告バーが出た", "未確認のページがあります。「最初の未確認へ」で戻るか「すべてスキップして進む」を押します。"],
        ["書き出し先がわからない", "書き出すときに自分でフォルダを選びます。デスクトップなど分かりやすい場所を選んでください。"],
        ["置換されない", "「辞書」タブで用語を登録し、「全ファイルに再適用」を押してください。本文も対象にするには「ヘッダのみに適用」のチェックを外します。"],
        ["別の PC でも同じ辞書を使いたい", "下の「辞書の共有」を参照してください。"],
        ["画面が真っ白／開かない", "アプリを閉じてもう一度起動します。直らないときは情報システム担当へ。"],
    ], widths=[4.6, 11.4])

    h(doc, "8. 辞書をほかの PC と共有する")
    numbered(doc, "「辞書」タブの **「⬇ JSON 書き出し」** を押し、辞書ファイル（dictionary.json）を保存します。")
    numbered(doc, "そのファイルを共有フォルダやメールで相手に渡します。")
    numbered(doc, "受け取った側は **「⬆ JSON 読み込み」** を押し、そのファイルを選ぶと辞書が取り込まれます。")

    out = DOCS / "PdfToSvg_操作手順書.docx"
    doc.save(str(out))
    print("saved", out.name)


if __name__ == "__main__":
    build_design()
    build_manual()
    print("done.")
