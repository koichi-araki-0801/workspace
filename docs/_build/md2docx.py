# -*- coding: utf-8 -*-
"""Markdown 原稿 → Word(.docx) 共通レンダラ。

`docs/<project>/src/*.md` を正典とし、本書のヘルパ群で .docx を生成する。フォント規約
（CLAUDE.md「Word ドキュメント」）は `BIZ UD` ファミリへ統一: 本文 `BIZ UDPGothic`
（ascii/eastAsia 両設定）、等幅は同ファミリ `BIZ UDGothic`。

文書種別で見た目を出し分ける（front-matter `style` / 無指定はファイル名・title から推定）:
  - `style: guide` 案3 カード型 … 操作手順書・配布運用手順書（番号バッジ・カード・全面グリッド）
  - `style: spec`  案2 テクニカル型 … 設計書（等幅アクセント番号・罫線中心・コード左罫線）

依存: python-docx（成果物生成のための一時依存。各ツール本体の requirements には足さない）。
外部 Markdown ライブラリは使わず、必要構文だけの行指向パーサを自前実装（オフライン依存追加を避けるため）。

対応構文の対応表は `docs/_build/README` 相当として本書冒頭コメントに集約する:
  - `#` / `##` / `###`            → `h(level)`（文書タイトルはフロントマター `title`/`subtitle`）
  - 空行区切りの段落             → `para`（`**強調**` と `` `等幅` `` のインライン対応）
  - `- `（字下げ）                → `bullet`（字下げ→level）
  - `1. `（番号付き）             → `numbered_card`（案3 level0）/ `numbered`（案2・ネスト）
  - ```` ``` ```` フェンス        → `code`（言語タグは案2 のラベルに使用）
  - GFM テーブル `| a | b |`      → `table`（案3=ACCENT ヘッダ＋ゼブラ / 案2=罫線中心）
  - `![caption](images/foo.png)` → `image`（枠付き 1 セル表）
  - `> [!WARN]/[!INFO]/[!NOTE]/[!TIP]/>` 引用 → `callout`（タグで地色・枠色・ラベルを選択）
  - `<!-- pagebreak -->`         → `page_break`

フロントマター（先頭 `---` ブロック・`key: value` のみの簡易 YAML）でメタを持つ:
  title / subtitle / out（出力 .docx 名）/ version / style（guide|spec）/ kicker（アイブロウ語）/
  images（画像基準ディレクトリの相対パス）/
  chapter_break（true で h1 章頭を強制改ページ。長文の設計書向けの任意オプション。先頭 h1 は除外）/
  toc（true で表紙直後に静的目次を描画。spec 向け。TOC フィールドは開くたび F9 更新が要るため
  自前描画とする。ページ番号は載らないが `chapter_break` 併用で章番号がナビゲーションになる）

spec 文書には加えてフッター（タイトル・版・`PAGE / NUMPAGES` フィールド）と、見出し段落への
`outlineLvl` 付与（Word ナビゲーションウィンドウ・PDF しおり用。見た目は不変）を自動で行う。

ページ送り制御（章のページまたぎ対策）は Word の段落プロパティで行う:
  見出し=`keepNext`+`keepLines`（孤立見出し防止）/ 全段落=`widowControl`（寡婦・孤児行抑制）/
  短いコード・callout=`keepLines`（ページ割れ防止。長大なら逆効果なので行数しきい値で打ち切り）。
"""
from __future__ import annotations

import pathlib
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── デザイントークン（手順書=案3 カード型 / 設計書=案2 テクニカル型 で共有）──
# フォントは `BIZ UD` ファミリへ統一（CLAUDE.md「Word ドキュメント」規約）。本文 `BIZ UDPGothic`
# （プロポーショナル）、等幅は同ファミリ `BIZ UDGothic`。両者は同じ BIZ UD なので混ざって見えない。
JP = "BIZ UDPGothic"      # 本文・見出し・表（統一）
MONO = "BIZ UDGothic"     # 等幅（コード・識別子・図のけた揃え）。同一 BIZ UD ファミリ
ACCENT = RGBColor(0x1F, 0x5C, 0x99)   # #1F5C99 見出しバー・表ヘッダ・番号バッジ・縦バー
INK = RGBColor(0x20, 0x24, 0x2C)      # #20242C 本文・見出し文字
MUTED = RGBColor(0x60, 0x68, 0x74)    # #606874 サブタイトル・キャプション
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# 網掛け fill / 罫線色（"RRGGBB" 文字列）。`_shade` / `_set_cell_border` 等の低レベルヘルパへ渡す。
FILL_ACCENT = "1F5C99"   # 見出しバー・表ヘッダ・番号バッジ
FILL_ZEBRA = "F5F8FB"    # 表の偶数行（案3）
FILL_HAIR = "DCE3EA"     # 細罫線・画像枠（案2）
FILL_CARD = "FAFBFC"     # カード地（案3 ステップ）
FILL_CARDBD = "D7DEE5"   # カード枠（案3）
FILL_CODE = "F4F6F8"     # コード網掛け（案2）
FILL_CODEBD = "C4CCD4"   # コード左罫線（案2）
FILL_RULE2 = "CDD5DD"    # 見出し下罫線（案2）
FILL_ROW2 = "E8ECF0"     # 行下罫線（案2 表）

# callout の `[!TAG]` → 背景色 / 外周罫線色。既定（タグ無し `>`）は INFO 系。
CALLOUT_FILL = {"WARN": "FDECEC", "INFO": "EAF2FB", "NOTE": "EAF2FB", "TIP": "FFF6E5"}
CALLOUT_BORDER = {"WARN": "E4A9A9", "INFO": "B9D3EE", "NOTE": "B9D3EE", "TIP": "F0D9A0"}
CALLOUT_DEFAULT = "INFO"   # タグ無し `>` 引用は INFO 扱い

# レンダリング中の文書スタイル。`render()` が原稿ごとに設定し、各ヘルパが分岐参照する。
#   "guide" = 案3 カード型（操作手順書・配布運用手順書）/ "spec" = 案2 テクニカル型（設計書）
_STYLE = "guide"

# 章のページまたぎ制御の状態（`render()` が原稿ごとにリセットし、`h()` が参照）。
_CHAPTER_BREAK = False   # front-matter `chapter_break` で h1 章頭の強制改ページを有効化
_H1_SEEN = False         # 先頭 h1 はタイトル直後なので改ページ対象から除外するためのフラグ

# コード/callout のページ割れ防止（`keepLines`）を適用する行数の上限。これを超える長大ブロックは
# 1 ページに収まらず `keepLines` が逆効果（手前を丸ごと次ページへ送る）になるため割れを許容する。
_KEEP_TOGETHER_MAXLINES = 25

# callout（案2）の `keepLines` 適用上限（文字数）。長文 callout（運用注記など）は 1 ページに
# 収まらないことがあり、`_KEEP_TOGETHER_MAXLINES` と同じ理由で割れを許容する。
_CALLOUT_KEEP_MAXCHARS = 400

# 案2 の表で偶数行ゼブラを自動適用するデータ行数のしきい値。短い表は「罫線中心・塗り最小」の
# 現行トーンを維持し、行を目で追いにくい長大な表だけに視線誘導を入れる。
_ZEBRA_MIN_ROWS = 8


# ── docx ヘルパ（`gen_docs.py` から移植）──
def _set_run_font(run, name=JP, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ── 低レベル罫線 / 網掛けヘルパ（python-docx に高水準 API が無いため OOXML を直接組む）──
def _shade(el, fill):
    """段落(`w:p`) or セル(`w:tc`)へ網掛けを設定。`el` は p要素 or tc要素。"""
    is_p = el.tag.endswith("}p")
    pr = el.get_or_add_pPr() if is_p else el.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _set_cell_border(cell, **edges):
    """セル罫線。`edges` 例: `top={'sz':4,'color':'DCE3EA'}`。`sz` は 1/8pt 単位。"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in edges:
            e = OxmlElement(f"w:{edge}")
            spec = edges[edge]
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(spec.get("sz", 4)))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), spec.get("color", FILL_HAIR))
            borders.append(e)
    tcPr.append(borders)


def _para_border(p, edge, sz, color, space):
    """段落 `p`（Paragraph ラッパ）の 1 辺に罫線を引く。`edge`∈{top,bottom,left}。`sz` は 1/8pt。"""
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pbdr.append(e)


def _para_bottom_border(p, sz=6, color=FILL_RULE2):
    """段落の下罫線（案2 見出しの下線）。"""
    _para_border(p, "bottom", sz, color, 4)


def _para_top_border(p, sz=16, color="20242C"):
    """段落の上罫線（案2 タイトルの全幅上線）。"""
    _para_border(p, "top", sz, color, 4)


def _para_left_border(p, sz=16, color=FILL_ACCENT, space=8):
    """段落の左罫線（callout の縦アクセント・案3 見出しバー等）。"""
    _para_border(p, "left", sz, color, space)


def _cant_split(row):
    """表行をページまたぎ禁止に（番号バッジ/カード/画像が割れないように）。"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _tbl_header(row):
    """表のヘッダ行をページまたぎで繰り返し表示する（`w:tblHeader`。ページを越える長大表向け）。"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def _outline_level(p, level):
    """見出し段落へアウトラインレベルを付与する（`w:outlineLvl`。h1→0）。見た目は不変のまま
    Word のナビゲーションウィンドウ・PDF 化時のしおりを効かせる。見出しを Heading スタイルに
    載せない方針（手動装飾を優先）のため、レベルだけを自前で埋める。"""
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level - 1))
    p._p.get_or_add_pPr().append(ol)


def _add_field_run(p, instr, size=8):
    """`PAGE` 等のフィールドを fldChar begin/instrText/end の 3 run で挿入する（python-docx に
    高水準 API が無いため OOXML を直接組む）。instrText の run にもフォントを設定し、Word が
    計算した数字が既定フォントへ落ちないようにする。"""
    rb = p.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    rb._element.append(fb)
    ri = p.add_run()
    _set_run_font(ri, name=MONO, size=size, color=MUTED)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    ri._element.append(it)
    rn = p.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    rn._element.append(fe)


def _run_shade(run, fill):
    """run 単位の網掛け（callout のタグ小箱など、文字の背景だけ塗る）。"""
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


# callout タグ → 先頭ラベル文字列（案3 は小箱、案2 は `ラベル —`）。
_CALLOUT_LABEL = {"WARN": "注意", "INFO": "INFO", "NOTE": "NOTE", "TIP": "ヒント"}


def new_doc() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = JP
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), JP)
    # 案2: 段落の寡婦・孤児行制御。最終 1 行だけが次ページ頭／先頭 1 行だけが前ページ末に残る
    # のを抑える。Normal に一括設定し全段落へ波及させる。
    normal.paragraph_format.widow_control = True
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
    return doc


def title(doc, text, subtitle=None, version=None, kicker=None):
    """表紙タイトル。案3=アクセント縦バー＋小ラベル＋28pt、案2=等幅アイブロウ＋上罫線＋右端に版。"""
    if _STYLE == "spec":
        # アイブロウ（等幅アクセント）＋全幅上罫線
        eb = doc.add_paragraph()
        _para_top_border(eb, sz=16, color="20242C")
        eb.paragraph_format.space_before = Pt(2)
        eb.paragraph_format.space_after = Pt(0)
        er = eb.add_run(kicker or "DESIGN SPEC")
        _set_run_font(er, name=MONO, size=9, bold=True, color=ACCENT)
        # タイトル（INK 28pt）。版は等幅 MUTED で右側へ続ける。
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        _set_run_font(r, size=28, bold=True, color=INK)
        if version:
            rv = p.add_run(f"   v{version}")
            _set_run_font(rv, name=MONO, size=11, color=MUTED)
        if subtitle:
            ps = doc.add_paragraph()
            ps.paragraph_format.space_after = Pt(8)
            rs = ps.add_run(subtitle)
            _set_run_font(rs, name=MONO, size=10, color=MUTED)
        return
    # 案3 カード型: 左アクセント縦バー＋小ラベル＋28pt タイトル＋サブ＋版
    if kicker:
        pk = doc.add_paragraph()
        _para_left_border(pk, sz=24, color=FILL_ACCENT, space=8)
        pk.paragraph_format.left_indent = Cm(0.25)
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(kicker)
        _set_run_font(rk, size=9, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    _para_left_border(p, sz=24, color=FILL_ACCENT, space=8)
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _set_run_font(r, size=28, bold=True, color=INK)
    tail = subtitle or ""
    if version:
        tail = f"{tail} ・ 版 {version}" if tail else f"版 {version}"
    if tail:
        ps = doc.add_paragraph()
        ps.paragraph_format.left_indent = Cm(0.25)
        ps.paragraph_format.space_after = Pt(8)
        rs = ps.add_run(tail)
        _set_run_font(rs, size=11, color=MUTED)


# 見出し先頭の章番号（`2.` `8.3` 等）を等幅アクセントで分離するための検出。
_HNUM = re.compile(r"^(\d+(?:\.\d+)*\.?)\s+(.*)$")


def toc(doc, entries):
    """静的目次（案2・front-matter `toc: true`）。`_collect_headings` で先読みした見出し一覧を
    自前で描画する。TOC フィールドは開くたび F9 更新が要り配布物に不向きなため使わない。
    h3 はノイズになるので載せない。"""
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(10)
    hp.paragraph_format.space_after = Pt(6)
    hr = hp.add_run("目次")
    _set_run_font(hr, size=15, bold=True, color=INK)
    _para_bottom_border(hp, sz=6, color=FILL_RULE2)
    for level, text in entries:
        if level >= 3:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        m = _HNUM.match(text)
        if level == 1:
            if m:
                rn = p.add_run(m.group(1) + " ")
                _set_run_font(rn, name=MONO, size=10.5, bold=True, color=ACCENT)
                text = m.group(2)
            _add_inline(p, text, size=10.5, base_bold=True, base_color=INK)
        else:
            p.paragraph_format.left_indent = Cm(0.5)
            if m:
                rn = p.add_run(m.group(1) + " ")
                _set_run_font(rn, name=MONO, size=9.5, color=MUTED)
                text = m.group(2)
            _add_inline(p, text, size=9.5, base_color=MUTED)
    page_break(doc)


def _spec_footer(doc, title_text, version=None):
    """案2 のフッター。左=タイトル＋版 / 右=`PAGE / NUMPAGES`（Word が表示時に自動再計算）。
    既定 Footer スタイルが持つ中央/右タブに 1 個目のタブが吸われないよう Normal へ切り替え、
    本文幅の右端タブを自前で置く。上辺の細罫線で本文の罫線トーンに揃える。"""
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.style = doc.styles["Normal"]
    _para_border(p, "top", 4, FILL_RULE2, 4)
    width = sec.page_width - sec.left_margin - sec.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)
    label = title_text + (f"　v{version}" if version else "")
    rl = p.add_run(label)
    _set_run_font(rl, size=8, color=MUTED)
    p.add_run("\t")
    _add_field_run(p, "PAGE")
    rs = p.add_run(" / ")
    _set_run_font(rs, name=MONO, size=8, color=MUTED)
    _add_field_run(p, "NUMPAGES")


def h(doc, text, level=1):
    """見出し。案3=h1 に左アクセント縦バー、案2=等幅アクセント番号＋h1 下罫線。番号は原文ママ温存。"""
    global _H1_SEEN
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(3)
    # 案1: 見出しは直後の本文と同ページに保ち（孤立見出し防止）、見出し自身も行分割しない。
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    # 案4: 章(h1)頭の強制改ページ（front-matter `chapter_break` 有効時のみ）。先頭 h1 は
    # タイトル直後なので除外し、2 つ目以降の h1 だけを次ページ送りにする。
    if level == 1 and _CHAPTER_BREAK:
        if _H1_SEEN:
            p.paragraph_format.page_break_before = True
        _H1_SEEN = True
    if _STYLE == "spec":
        _outline_level(p, level)
        size = {1: 15, 2: 13, 3: 12}.get(level, 11)
        m = _HNUM.match(text)
        if m:
            rn = p.add_run(m.group(1) + " ")
            _set_run_font(rn, name=MONO, size=size, bold=True, color=ACCENT)
            _add_inline(p, m.group(2), size=size, base_bold=True, base_color=INK)
        else:
            _add_inline(p, text, size=size, base_bold=True, base_color=INK)
        if level == 1:
            _para_bottom_border(p, sz=6, color=FILL_RULE2)
        return p
    # 案3 カード型
    size = {1: 16, 2: 13, 3: 12}.get(level, 11)
    if level == 1:
        _para_left_border(p, sz=24, color=FILL_ACCENT, space=8)
        p.paragraph_format.left_indent = Cm(0.2)
    else:
        p.paragraph_format.left_indent = Cm(0.15)
    _add_inline(p, text, size=size, base_bold=True, base_color=INK)
    return p


def para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _add_inline(p, text, size=size)
    return p


# `**強調**` と `` `等幅` `` を runs へ分解するインラインパーサ。
_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _add_inline(p, text, size=10.5, base_bold=False, base_color=None):
    """`text` を解析し、強調/等幅を反映した runs を段落 `p` へ追加する。"""
    parts = _INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            r = p.add_run(part[2:-2])
            _set_run_font(r, size=size, bold=True, color=base_color)
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            r = p.add_run(part[1:-1])
            _set_run_font(r, name=MONO, size=size - 0.5, bold=base_bold, color=base_color)
        else:
            r = p.add_run(part)
            _set_run_font(r, size=size, bold=base_bold, color=base_color)


def bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    _add_inline(p, text, size=size)
    return p


def numbered(doc, text, n, level=0, size=10.5):
    """番号付きステップ（案2・案3 ネスト用）。`List Number` は使わず等幅アクセント番号のぶら下げ。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    rn = p.add_run(f"{n}. ")
    _set_run_font(rn, name=MONO, size=size, bold=True, color=ACCENT)
    _add_inline(p, text, size=size)
    return p


def numbered_card(doc, text, n):
    """案3 カード型のステップ。2列1行表（番号セル=アクセント白太字 / 本文セル=カード地）。"""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(0.95)   # 番号セル
    t.columns[1].width = Cm(15.0)   # 本文セル
    num_cell, body_cell = t.rows[0].cells
    num_cell.width = Cm(0.95)
    body_cell.width = Cm(15.0)
    # 番号セル: アクセント塗り・白太字・縦中央
    _shade(num_cell._tc, FILL_ACCENT)
    num_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rp = num_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rp.add_run(str(n))
    _set_run_font(r, size=11, bold=True, color=WHITE)
    # 本文セル
    bp = body_cell.paragraphs[0]
    _add_inline(bp, text, size=10.5, base_color=INK)
    # カード枠（両セル外周）＋本文セル地
    for c in (num_cell, body_cell):
        _set_cell_border(
            c, top={"sz": 4, "color": FILL_CARDBD}, bottom={"sz": 4, "color": FILL_CARDBD},
            left={"sz": 4, "color": FILL_CARDBD}, right={"sz": 4, "color": FILL_CARDBD})
    _shade(body_cell._tc, FILL_CARD)
    _cant_split(t.rows[0])
    return t


def code(doc, text, size=8.8, lang=None):
    """コードブロック。案2=`FILL_CODE` 網掛け＋左2pt 罫線＋言語ラベル。案3=淡い網掛けのみ。"""
    if _STYLE == "spec" and lang:
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Cm(0.3)
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(lang.upper())
        _set_run_font(lr, name=MONO, size=8, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if (_STYLE == "spec" and lang) else 3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    if _STYLE == "spec":
        _shade(p._p, FILL_CODE)
        _para_left_border(p, sz=16, color=FILL_CODEBD, space=6)
        size = 9.0
    else:
        _shade(p._p, "F4F6F8")
    # 案3: 短いコードはページ割れ防止。長大なら 1 ページに収まらず逆効果なので割れを許容する。
    if text.count("\n") + 1 <= _KEEP_TOGETHER_MAXLINES:
        p.paragraph_format.keep_together = True
    r = p.add_run(text)
    _set_run_font(r, name=MONO, size=size)


# 表セル外周をカード枠（案3）で囲むときの 4 辺指定。
_GRID4 = dict(top={"sz": 4, "color": FILL_CARDBD}, bottom={"sz": 4, "color": FILL_CARDBD},
              left={"sz": 4, "color": FILL_CARDBD}, right={"sz": 4, "color": FILL_CARDBD})


def table(doc, header, rows, widths=None):
    """GFM テーブル。案3=全面グリッド＋ACCENT ヘッダ＋ゼブラ、案2=罫線中心・塗りなし。"""
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    spec = _STYLE == "spec"
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        if spec:
            _add_inline(hdr[i].paragraphs[0], htext, size=10, base_bold=True, base_color=ACCENT)
            _set_cell_border(hdr[i], bottom={"sz": 12, "color": "20242C"})
        else:
            _add_inline(hdr[i].paragraphs[0], htext, size=9.5, base_bold=True, base_color=WHITE)
    if spec:
        # 長大表対策: ヘッダ行はページまたぎで繰り返し、淡塗りでデータ行との見分けを付ける。
        _tbl_header(t.rows[0])
        for c in hdr:
            _shade(c._tc, FILL_CODE)
    else:
        # §6-A: style 後ではなく塗りで直接アクセントヘッダにする（淡色潰れ回避）。
        for c in hdr:
            _shade(c._tc, FILL_ACCENT)
            _set_cell_border(c, **_GRID4)
    zebra = spec and len(rows) >= _ZEBRA_MIN_ROWS
    for ridx, row in enumerate(rows):
        tr = t.add_row()
        cells = tr.cells
        for i, val in enumerate(row):
            cells[i].text = ""
            _add_inline(cells[i].paragraphs[0], str(val), size=9.5)
        if spec:
            _cant_split(tr)   # 行は 1〜2 行文なので分割禁止の副作用はない
            for c in cells:
                _set_cell_border(c, bottom={"sz": 8, "color": FILL_ROW2})
            if zebra and ridx % 2 == 1:   # データ 1 件目=奇数 → 偶数行に zebra（案3 と同じ位相）
                for c in cells:
                    _shade(c._tc, FILL_ZEBRA)
        else:
            if ridx % 2 == 1:   # データ 1 件目=奇数 → 偶数行に zebra
                for c in cells:
                    _shade(c._tc, FILL_ZEBRA)
            for c in cells:
                _set_cell_border(c, **_GRID4)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def image(doc, img_dir: pathlib.Path, name, width_cm=16.0, caption=None, figno=None):
    """`img_dir/name` を 1 セル表の枠付きで中央挿入。欠落時はプレースホルダ（警告は呼び元が集計）。"""
    path = img_dir / name
    if not path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"［画像未取得: {name}］")
        _set_run_font(r, size=9.5, italic=True, color=MUTED)
        return False
    border = FILL_HAIR if _STYLE == "spec" else FILL_CARDBD
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _set_cell_border(cell, top={"sz": 8, "color": border}, bottom={"sz": 8, "color": border},
                     left={"sz": 8, "color": border}, right={"sz": 8, "color": border})
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    _cant_split(t.rows[0])
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        if _STYLE == "spec":
            lead = f"fig.{figno} " if figno else "fig. "
            rl = c.add_run(lead)
            _set_run_font(rl, name=MONO, size=9, color=MUTED)
        r = c.add_run(caption)
        _set_run_font(r, size=9, italic=(_STYLE != "spec"), color=MUTED)
    return True


def callout(doc, text, tag=None):
    """補足ボックス。案3=全面網掛けカード＋外周罫線＋タグ小箱、案2=左2pt アクセント罫線＋等幅ラベル。"""
    tag = (tag or CALLOUT_DEFAULT).upper()
    if tag not in CALLOUT_FILL:
        tag = "INFO"
    label = _CALLOUT_LABEL.get(tag, tag)
    if _STYLE == "spec":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.2)
        # 補足ボックスは短いのが通例なのでページ割れを防ぐ（左アクセント罫線が分断されない）。
        # ただし長文（運用注記など）は 1 ページに収まらず逆効果になるため文字数で打ち切る。
        if len(text) <= _CALLOUT_KEEP_MAXCHARS:
            p.paragraph_format.keep_together = True
        _para_left_border(p, sz=16, color=FILL_ACCENT, space=8)
        lr = p.add_run(f"{label} — ")
        _set_run_font(lr, name=MONO, size=9.5, bold=True, color=ACCENT)
        _add_inline(p, text, size=10)
        return
    # 案3: 全面網掛けカード（1 セル表）＋外周罫線＋先頭タグ小箱
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell._tc, CALLOUT_FILL[tag])
    border = CALLOUT_BORDER[tag]
    _set_cell_border(cell, top={"sz": 8, "color": border}, bottom={"sz": 8, "color": border},
                     left={"sz": 8, "color": border}, right={"sz": 8, "color": border})
    p = cell.paragraphs[0]
    chip = p.add_run(f" {label} ")
    _set_run_font(chip, size=9, bold=True, color=WHITE)
    _run_shade(chip, FILL_ACCENT)
    p.add_run("  ")
    _add_inline(p, text, size=10, base_color=INK)
    _cant_split(t.rows[0])


def page_break(doc):
    doc.add_page_break()


# ── front-matter + Markdown パーサ ──
def _parse_frontmatter(text: str):
    """先頭 `---` ブロックを `key: value` の辞書として取り出し、本文を返す。"""
    meta = {}
    lines = text.splitlines()
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            line = lines[i]
            if ":" in line:
                k, _, v = line.partition(":")
                meta[k.strip()] = v.strip().strip('"').strip("'")
            i += 1
        body = "\n".join(lines[i + 1:]) if i < len(lines) else ""
        return meta, body
    return meta, text


def _meta_flag(meta, key):
    """front-matter の真偽値オプション（`true`/`1`/`yes`/`on` を真とする）。"""
    return str(meta.get(key, "")).strip().lower() in ("true", "1", "yes", "on")


def _collect_headings(body):
    """本文から (level, text) を先読み収集する（目次描画用）。コードフェンス内の `#` 行を
    見出しと誤認しないようフェンス状態を追跡する。level は本パーサと同じく 3 で頭打ち。"""
    out = []
    in_fence = False
    for line in body.splitlines():
        s = line.strip()
        if s.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            out.append((min(len(m.group(1)), 3), m.group(2).strip()))
    return out


_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")
_IMG = re.compile(r"^!\[(?P<cap>.*?)\]\((?P<src>.*?)\)\s*$")
_CALLOUT = re.compile(r"^>\s?(?:\[!(?P<tag>[A-Z]+)\]\s*)?(?P<body>.*)$")


def _split_row(line: str):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def render(md_path, out_path, img_dir, warnings=None):
    """Markdown 原稿 `md_path` を `out_path` の .docx へ描画する。

    `img_dir` は画像参照の基準ディレクトリ。`warnings` は欠落画像などを積む list（任意）。
    戻り値はフロントマター辞書。
    """
    md_path = pathlib.Path(md_path)
    out_path = pathlib.Path(out_path)
    img_dir = pathlib.Path(img_dir)
    warnings = warnings if warnings is not None else []

    meta, body = _parse_frontmatter(md_path.read_text(encoding="utf-8"))
    images_override = meta.get("images")
    if images_override:
        img_dir = (md_path.parent / images_override).resolve()

    # 文書スタイル決定（front-matter `style` 優先 / 未指定はファイル名・title から推定）。
    global _STYLE, _CHAPTER_BREAK, _H1_SEEN
    style = (meta.get("style") or "").strip().lower()
    if style not in ("guide", "spec"):
        hay = (meta.get("out", "") + meta.get("title", "") + md_path.stem)
        style = "spec" if "設計" in hay else "guide"
    _STYLE = style

    # 案4: 章頭改ページの可否を front-matter から取得し、先頭 h1 判定フラグを原稿ごとにリセット。
    _CHAPTER_BREAK = _meta_flag(meta, "chapter_break")
    _H1_SEEN = False

    doc = new_doc()
    if meta.get("title"):
        # 版は title 内（案3=サブに続け / 案2=右端）で表示。別段落は出さない。
        title(doc, meta["title"], meta.get("subtitle"),
              version=meta.get("version"), kicker=meta.get("kicker"))
    # 案2 の長文ナビゲーション: 静的目次（`toc: true` 時）とページ番号入りフッター。
    if _STYLE == "spec" and _meta_flag(meta, "toc"):
        toc(doc, _collect_headings(body))
    if _STYLE == "spec" and meta.get("title"):
        _spec_footer(doc, meta["title"], meta.get("version"))

    lines = body.splitlines()
    i = 0
    n = len(lines)
    fig_no = 0
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # ページ区切り
        if stripped == "<!-- pagebreak -->":
            page_break(doc)
            i += 1
            continue

        # コードフェンス
        if stripped.startswith("```"):
            lang = stripped[3:].strip() or None
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # 終了フェンスを飛ばす
            code(doc, "\n".join(buf), lang=lang)
            continue

        # 見出し
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            h(doc, m.group(2).strip(), level=level)
            i += 1
            continue

        # 画像（単独行）
        mi = _IMG.match(stripped)
        if mi:
            name = pathlib.Path(mi.group("src")).name
            fig_no += 1
            ok = image(doc, img_dir, name, caption=(mi.group("cap") or None), figno=fig_no)
            if not ok:
                warnings.append(f"{md_path.name}: 画像未取得 {name}")
            i += 1
            continue

        # callout（引用）
        if stripped.startswith(">"):
            mc = _CALLOUT.match(stripped)
            tag = (mc.group("tag") if mc else None)
            parts = [mc.group("body") if mc else stripped[1:].strip()]
            i += 1
            while i < n and lines[i].strip().startswith(">"):
                mc2 = _CALLOUT.match(lines[i].strip())
                parts.append(mc2.group("body") if mc2 else lines[i].strip()[1:].strip())
                i += 1
            callout(doc, " ".join(p for p in parts if p), tag=tag)
            continue

        # テーブル（ヘッダ行 + 区切り行）
        if stripped.startswith("|") and i + 1 < n and _TABLE_SEP.match(lines[i + 1]):
            header = _split_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            table(doc, header, rows)
            continue

        # 箇条書き / 番号付き
        mb = re.match(r"^(?P<indent>\s*)(?P<marker>[-*]|\d+\.)\s+(?P<text>.*)$", line)
        if mb:
            level = len(mb.group("indent")) // 2
            txt = mb.group("text").strip()
            marker = mb.group("marker")
            if marker.endswith("."):
                num = int(marker[:-1])   # 原文の採番を温存
                if _STYLE == "guide" and level == 0:
                    numbered_card(doc, txt, num)   # 案3: 番号バッジ付きカード
                else:
                    numbered(doc, txt, num, level=level)
            else:
                bullet(doc, txt, level=level)
            i += 1
            continue

        # 段落（空行 / 別ブロック開始まで集約）
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith(">")
                    or nxt.startswith("```") or nxt.startswith("|")
                    or _IMG.match(nxt) or nxt == "<!-- pagebreak -->"
                    or re.match(r"^([-*]|\d+\.)\s+", nxt)):
                break
            buf.append(nxt)
            i += 1
        para(doc, " ".join(buf))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return meta
